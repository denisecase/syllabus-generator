"""src/syllabus_generator/generator.py.

Generate syllabus documents from a Word template.
"""

from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.document import Document as DocumentType
from docx.table import Table, _Cell
from docx.text.paragraph import Paragraph

from syllabus_generator.constants import (
    GRADING_PLACEHOLDER,
    OUTPUTS_DIR,
    SCHEDULE_PLACEHOLDER,
)
from syllabus_generator.errors import TemplateError
from syllabus_generator.models import CourseData
from syllabus_generator.naming import build_output_filename


def _cell_text(cell: _Cell) -> str:
    """Return combined text for a cell."""
    return "\n".join(paragraph.text for paragraph in cell.paragraphs).strip()


def _set_cell_text(cell: _Cell, text: str) -> None:
    """Replace all content in a cell with plain text."""
    cell.text = text


def _find_table_with_placeholder(
    document: DocumentType, placeholder: str
) -> Table | None:
    """Find the first table containing the given placeholder."""
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if placeholder in _cell_text(cell):
                    return table
    return None


def _find_row_index_with_placeholder(table: Table, placeholder: str) -> int | None:
    """Find the row index containing the given placeholder."""
    for idx, row in enumerate(table.rows):
        for cell in row.cells:
            if placeholder in _cell_text(cell):
                return idx
    return None


def _replace_learning_outcomes_table(
    document: DocumentType, course: CourseData
) -> None:
    """Replace the learning outcomes marker row with one row per outcome."""
    desc_placeholder = "{{learning_outcomes_description}}"

    table = _find_table_with_placeholder(document, desc_placeholder)
    if table is None:
        return

    row_index = _find_row_index_with_placeholder(table, desc_placeholder)
    if row_index is None:
        return

    template_row = table.rows[row_index]
    template_tr = template_row._tr
    parent_tbl = template_tr.getparent()

    for item in course.learning_outcomes:
        new_tr = deepcopy(template_tr)
        parent_tbl.insert(parent_tbl.index(template_tr), new_tr)

        new_row = table.rows[row_index]
        left_text = f"{item.id}. {item.description}"
        right_text = ", ".join(item.assessment)

        _set_cell_text(new_row.cells[0], left_text)
        _set_cell_text(new_row.cells[1], right_text)

        row_index += 1

    parent_tbl.remove(template_tr)


def _replace_in_paragraph(paragraph: Paragraph, replacements: dict[str, str]) -> None:
    """Replace placeholders in a paragraph."""
    if not paragraph.text:
        return

    updated = paragraph.text
    changed = False

    for key, value in replacements.items():
        if key in updated:
            updated = updated.replace(key, value)
            changed = True

    if changed:
        paragraph.text = updated


def _replace_in_table(table: Table, replacements: dict[str, str]) -> None:
    """Replace placeholders in a table."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                _replace_in_paragraph(paragraph, replacements)
            for nested_table in cell.tables:
                _replace_in_table(nested_table, replacements)


def _replace_everywhere(document: "DocumentType", replacements: dict[str, str]) -> None:
    """Replace placeholders in paragraphs and tables."""
    for paragraph in document.paragraphs:
        _replace_in_paragraph(paragraph, replacements)

    for table in document.tables:
        _replace_in_table(table, replacements)


def _replace_grading_table(document: "DocumentType", course: CourseData) -> None:
    comp_placeholder = "{{grading_component}}"

    table = _find_table_with_placeholder(document, comp_placeholder)
    if table is None:
        return

    row_index = _find_row_index_with_placeholder(table, comp_placeholder)
    if row_index is None:
        return

    template_row = table.rows[row_index]
    template_tr = template_row._tr
    parent_tbl = template_tr.getparent()

    for item in course.grading:
        new_tr = deepcopy(template_tr)
        parent_tbl.insert(parent_tbl.index(template_tr), new_tr)

        new_row = table.rows[row_index]

        _set_cell_text(new_row.cells[0], item.component)
        _set_cell_text(new_row.cells[1], f"{item.weight}%")

        row_index += 1

    parent_tbl.remove(template_tr)


def _format_grading(course: CourseData) -> str:
    """Format grading section for block replacement."""
    lines = [f"{item.component}: {item.weight}%" for item in course.grading]
    return "\n".join(lines)


def _format_schedule(course: CourseData) -> str:
    """Format schedule section for block replacement."""
    lines = [
        f"Week {item.week}: {item.topic} | Deliverable: {item.deliverable}"
        for item in course.schedule
    ]
    return "\n".join(lines)


def _build_module_replacements(course: CourseData) -> dict[str, str]:
    """Build placeholder replacements for module0 ... module7."""
    replacements: dict[str, str] = {}

    for i in range(8):
        key = f"module{i}"
        replacements[f"{{{{{key}}}}}"] = str(course.modules.get(key, "")).strip()

    return replacements


def build_replacements(course: CourseData) -> dict[str, str]:
    """Build placeholder replacement values."""
    replacements = {
        "{{course_prefix}}": course.course_prefix,
        "{{course_number}}": course.course_number,
        "{{course_sections}}": course.course_sections,
        "{{course_year}}": course.course_year,
        "{{course_yy}}": course.course_yy,
        "{{course_season}}": course.course_season,
        "{{course_block}}": course.course_block,
        "{{course_name_line1}}": course.course_name_line1,
        "{{course_name_line2}}": course.course_name_line2,
        "{{course_name_full}}": course.course_name_full,
        "{{course_short_name}}": course.course_short_name,
        "{{materials_textbook}}": course.materials_textbook,
        "{{materials_other}}": course.materials_other,
        "{{placement}}": course.placement,
        "{{prereqs}}": course.prereqs,
        "{{credits}}": course.credits,
        "{{description}}": course.description,
        "{{rationale}}": course.rationale,
        "{{instructor_name}}": course.instructor.name,
        "{{instructor_last_name}}": course.instructor.last_name,
        "{{instructor_email}}": course.instructor.email,
        "{{instructor_office_hours}}": course.instructor.office_hours,
        "{{instructor_aside}}": course.instructor.aside,
        GRADING_PLACEHOLDER: _format_grading(course),
        SCHEDULE_PLACEHOLDER: _format_schedule(course),
    }
    replacements.update(_build_module_replacements(course))
    return replacements


def generate_syllabus(
    course: CourseData,
    template_path: Path,
    output_dir: Path = OUTPUTS_DIR,
) -> Path:
    """Generate one syllabus document and return the output path."""
    try:
        document: DocumentType = Document(str(template_path))
    except Exception as exc:
        raise TemplateError(f"Could not open Word template: {template_path}") from exc

    _replace_learning_outcomes_table(document, course)
    _replace_grading_table(document, course)

    replacements = build_replacements(course)
    _replace_everywhere(document, replacements)

    output_dir.mkdir(parents=True, exist_ok=True)
    output_path = output_dir / build_output_filename(course)

    try:
        document.save(str(output_path))
    except Exception as exc:
        raise TemplateError(f"Could not save output document: {output_path}") from exc

    return output_path
